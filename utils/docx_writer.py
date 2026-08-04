"""Write Word (.docx) release notes from agent outputs."""
from __future__ import annotations

from pathlib import Path

from docx import Document


def write_release_notes(
    path: Path,
    title: str,
    requirement: str,
    sections: dict[str, str],
) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_heading("Original Requirement", level=1)
    doc.add_paragraph(requirement)
    for name, content in sections.items():
        doc.add_heading(name, level=1)
        # Keep Word docs readable
        for para in str(content).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.save(str(path))
    return path
